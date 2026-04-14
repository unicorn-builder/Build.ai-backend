"""
gen_rapport_docx_en.py — Executive Report as Word (.docx) — English
Tijan AI — same data as gen_rapport_docx.py, English labels/headers
"""
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


VERT = RGBColor(0x43, 0xA9, 0x56)
GRIS = RGBColor(0x55, 0x55, 0x55)


def _styled_table(doc, headers, rows, col_widths=None):
    """Add a table with green header."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shading = cell._element.get_or_add_tcPr()
        shading_elm = shading.makeelement(qn('w:shd'), {
            qn('w:fill'): '43A956', qn('w:val'): 'clear'
        })
        shading.append(shading_elm)

    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    return table


def _fmt(n):
    """Format number with thousands separator."""
    if isinstance(n, (int, float)):
        return f'{n:,.0f}'
    return str(n)


def generer_rapport_executif_docx(rs, rm, params: dict) -> bytes:
    """Generate Executive Report as Word — English. Returns docx bytes."""
    doc = Document()
    d = rs.params
    boq_s = rs.boq
    boq_m = rm.boq
    e = rm.edge

    # Title
    title = doc.add_heading(d.nom, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = VERT

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'Executive Summary Report — {d.ville} — R+{d.nb_niveaux-1}')
    run.font.size = Pt(12)
    run.font.color.rgb = GRIS

    doc.add_paragraph(
        'This document is intended for the project owner. '
        'It presents the key points of the project and overall budget estimate.'
    )

    # Section 1 — Project Sheet
    doc.add_heading('1. PROJECT DATA', level=1)
    fiche_data = [
        ('Project', d.nom, 'Location', f'{d.ville}, {d.pays}'),
        ('Use', d.usage.value.capitalize(), 'Height', f'R+{d.nb_niveaux-1} ({d.nb_niveaux} levels)'),
        ('Built area', f'{_fmt(boq_s.surface_batie_m2)} m²', 'Habitable area', f'{_fmt(boq_s.surface_habitable_m2)} m²'),
        ('Dwellings', str(rm.nb_logements), 'Occupants', str(rm.nb_personnes)),
        ('Concrete', rs.classe_beton, 'Foundation', rs.fondation.type.value),
        ('EDGE certification', '✓ Certifiable' if e.certifiable else '✗ Not certifiable', 'EC2 compliance', rs.analyse.conformite_ec2),
    ]
    _styled_table(doc,
        ['Parameter', 'Value', 'Parameter', 'Value'],
        fiche_data,
        col_widths=[4, 4, 4, 4],
    )

    # Section 2 — Budget
    doc.add_heading('2. OVERALL BUDGET ESTIMATE', level=1)
    doc.add_paragraph('Estimation ±15% — Local market unit prices 2026').italic = True

    total_bas = boq_s.total_bas_fcfa + boq_m.total_basic_fcfa
    total_haut = boq_s.total_haut_fcfa + boq_m.total_hend_fcfa
    pct_s = boq_s.total_bas_fcfa / total_bas * 100 if total_bas else 0
    pct_m = boq_m.total_basic_fcfa / total_bas * 100 if total_bas else 0

    budget_data = [
        ('Structure (main structure)', _fmt(boq_s.total_bas_fcfa), _fmt(boq_s.total_haut_fcfa), f'{pct_s:.0f}%', 'Concrete + steel + foundations'),
        ('MEP — Basic', _fmt(boq_m.total_basic_fcfa), _fmt(boq_m.total_hend_fcfa), f'{pct_m:.0f}%', 'Electrical, plumbing, HVAC, lifts, safety'),
        ('TOTAL MAIN STRUCTURE', _fmt(total_bas), _fmt(total_haut), '100%', 'Excludes finishes, site works'),
        ('Finishes (est. 35%)', _fmt(int(total_bas*0.35)), _fmt(int(total_haut*0.35)), '~35%', 'Tiling, painting, etc.'),
        ('TOTAL ESTIMATED COST', _fmt(int(total_bas*1.35)), _fmt(int(total_haut*1.35)), '', f'{int(total_bas/boq_s.surface_batie_m2):,} FCFA/m²'),
    ]
    _styled_table(doc,
        ['Trade', 'Low amount (FCFA)', 'High amount', '% Total', 'Note'],
        budget_data,
        col_widths=[4, 3.5, 3.5, 2, 4],
    )

    # Section 3 — EDGE
    doc.add_heading('3. ENVIRONMENTAL PERFORMANCE (EDGE IFC)', level=1)
    edge_data = [
        ('Energy savings', f'{e.economie_energie_pct:.0f}%', '≥ 20%', '✓' if e.economie_energie_pct >= 20 else '✗'),
        ('Water savings', f'{e.economie_eau_pct:.0f}%', '≥ 20%', '✓' if e.economie_eau_pct >= 20 else '✗'),
        ('Material savings', f'{e.economie_materiaux_pct:.0f}%', '≥ 20%', '✓' if e.economie_materiaux_pct >= 20 else '✗'),
        ('VERDICT', e.niveau_certification, '3/3 pillars', '✓ CERTIFIABLE' if e.certifiable else '✗ NOT CERTIFIABLE'),
    ]
    _styled_table(doc, ['Pillar', 'Score', 'Threshold', 'Status'], edge_data, col_widths=[5, 3, 3, 5])

    # Section 4 — Analysis/Insights
    doc.add_heading('4. TECHNICAL SUMMARY AND RECOMMENDATIONS', level=1)

    # Strengths
    if rs.analyse.points_forts:
        doc.add_heading('Project strengths', level=2)
        for point in rs.analyse.points_forts:
            p = doc.add_paragraph(point, style='List Bullet')
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

    # Alerts
    if rs.analyse.alertes:
        doc.add_heading('Points to monitor', level=2)
        for alerte in rs.analyse.alertes:
            p = doc.add_paragraph(alerte, style='List Bullet')
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            for run in p.runs:
                run.font.color.rgb = RGBColor(0xFF, 0x99, 0x00)

    # Recommendations
    if rs.analyse.recommandations:
        doc.add_heading('Recommendations', level=2)
        for rec in rs.analyse.recommandations[:5]:
            p = doc.add_paragraph(rec, style='List Bullet')
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)

    # Disclaimer
    doc.add_paragraph()
    disclaimer = doc.add_paragraph(
        'This report is an indicative pre-study document (±15%). '
        'It does not replace technical studies by a licensed engineering firm, '
        'whose intervention is legally required before any construction work begins.'
    )
    disclaimer.paragraph_format.left_indent = Cm(0.5)
    disclaimer.paragraph_format.right_indent = Cm(0.5)
    for run in disclaimer.runs:
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = GRIS

    # Footer
    doc.add_paragraph('\n\nDocument generated by Tijan AI — tijan.ai').italic = True

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
