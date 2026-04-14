"""
gen_note_docx_en.py — Structural Calculation Note as Word (.docx) — English
Tijan AI — same data as gen_note_docx.py, English labels/headers
"""
import io
import logging
from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)

# Tijan theme colors
VERT = RGBColor(0x43, 0xA9, 0x56)
VERT_LIGHT = RGBColor(0xE8, 0xF5, 0xEA)
NAVY = RGBColor(0x1B, 0x2A, 0x4A)
ORANGE = RGBColor(0xFF, 0xB3, 0x0F)
ORANGE_LT = RGBColor(0xFF, 0xE5, 0xCC)
NOIR = RGBColor(0x11, 0x11, 0x11)
GRIS = RGBColor(0x55, 0x55, 0x55)
BLANC = RGBColor(0xFF, 0xFF, 0xFF)


def _format_number(val, decimals=0, unit=''):
    """Format number with comma thousands separator."""
    if val is None:
        return '—'
    if isinstance(val, str):
        return val
    if decimals == 0:
        return f'{int(val):,}'.replace(',', ' ') + (f' {unit}' if unit else '')
    return f'{val:,.{decimals}f}'.replace(',', ' ') + (f' {unit}' if unit else '')


def _format_fcfa(val):
    """Format value as FCFA currency."""
    if val is None or val == 0:
        return '—'
    return f'{int(val):,} FCFA'.replace(',', ' ')


def _set_cell_background(cell, color_hex):
    """Set cell background color."""
    shading = cell._element.get_or_add_tcPr()
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color_hex)
    shading.append(shading_elm)


def _add_styled_table(doc, headers, rows, col_widths=None, header_color='43A956', zebra=False):
    """Add a professionally styled table with header row shading."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set column widths if provided
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = BLANC
        _set_cell_background(cell, header_color)

    # Data rows with optional zebra striping
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            cell.text = str(val)
            p = cell.paragraphs[0]
            run = p.runs[0] if p.runs else p.add_run(str(val))
            run.font.size = Pt(8)

            # Zebra striping
            if zebra and r_idx % 2 == 0:
                _set_cell_background(cell, 'F5F5F5')

    return table


def generer(donnees, resultats, **kw) -> bytes:
    """
    Generate Structural Calculation Note as Word (.docx) — English.

    Args:
        donnees: DonneesProjet object
        resultats: ResultatsStructure object
        **kw: additional keyword arguments (ignored)

    Returns:
        bytes: DOCX document content
    """
    doc = Document()

    # Extract data
    d = donnees
    rs = resultats
    boq = rs.boq
    ana = rs.analyse
    surf_batie_estimee = not hasattr(d, 'surface_batie_plans') or not d.surface_batie_plans

    # ── HEADER ──────────────────────────────────────────────────
    title = doc.add_heading(d.nom, level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.color.rgb = VERT
        run.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f'{d.ville} — {d.usage.value.capitalize()} R+{d.nb_niveaux-1} ({d.nb_niveaux} levels)')
    run.font.size = Pt(12)
    run.font.color.rgb = GRIS

    disc = doc.add_paragraph('Indicative calculations ±15% — Must be verified by a licensed structural engineer.')
    disc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in disc.runs:
        run.font.italic = True
        run.font.color.rgb = GRIS

    doc.add_paragraph()  # Spacer

    # ── SECTION 1: PROJECT DATA ────────────────────────────
    doc.add_heading('1. PROJECT DATA', level=1)

    note_surf = ' *' if surf_batie_estimee else ''
    fiche_data = [
        ['Project', d.nom, 'Location', d.ville],
        ['Use', d.usage.value.capitalize(), 'Levels', f'R+{d.nb_niveaux-1} ({d.nb_niveaux})'],
        [f'Built area{note_surf}', _format_number(boq.surface_batie_m2, 0, 'm²'), 'Habitable area', _format_number(boq.surface_habitable_m2, 0, 'm²')],
        ['Spans', f'{d.portee_min_m}–{d.portee_max_m} m', 'Bays', f'{d.nb_travees_x}×{d.nb_travees_y}'],
        ['Concrete', f'{rs.classe_beton} — fck={rs.fck_MPa:.0f} MPa', 'Steel', f'{rs.classe_acier} — fyk={rs.fyk_MPa:.0f} MPa'],
        ['Bearing pressure', f'{rs.pression_sol_MPa} MPa', 'Distance to sea', f'{rs.distance_mer_km:.1f} km'],
        ['Loads G / Q', f'{rs.charge_G_kNm2} / {rs.charge_Q_kNm2} kN/m²', 'Seismic zone', f'Zone {rs.zone_sismique} — ag={rs.sismique.ag_g}g'],
    ]
    _add_styled_table(doc, ['Parameter', 'Value', 'Parameter', 'Value'], fiche_data,
                      col_widths=[4.2, 4.2, 4.2, 4.2], header_color='43A956')

    if surf_batie_estimee:
        note_p = doc.add_paragraph('* Built area estimated (footprint × levels) — to be confirmed with final drawings.')
        for run in note_p.runs:
            run.font.size = Pt(8)
            run.font.italic = True

    justif_p = doc.add_paragraph(f'ℹ {ana.justification_materiaux}')
    justif_p.paragraph_format.left_indent = Cm(0.5)
    for run in justif_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = GRIS

    doc.add_paragraph()  # Spacer

    # ── SECTION 2: DESIGN ASSUMPTIONS AND CODES ─────────────────────────
    doc.add_heading('2. DESIGN ASSUMPTIONS AND APPLICABLE CODES', level=1)

    hyp_data = [
        ['Reinforced concrete', 'Eurocode 2 — NF EN 1992-1-1', f'γc=1.5 — fcd={rs.fck_MPa/1.5:.1f} MPa'],
        ['Seismic', 'Eurocode 8 — NF EN 1998-1', f'Zone {rs.zone_sismique} — ag={rs.sismique.ag_g}g — DCL'],
        ['Dead load G', 'EC1 — NF EN 1991-1-1', f'{rs.charge_G_kNm2} kN/m² ({d.usage.value})'],
        ['Live load Q', 'EC1 — NF EN 1991-1-1', f'{rs.charge_Q_kNm2} kN/m² ({d.usage.value})'],
        ['ULS combination', '1.35G + 1.5Q', f'{1.35*rs.charge_G_kNm2+1.5*rs.charge_Q_kNm2:.1f} kN/m²'],
        ['Foundations', 'EC7 + DTU 13.2', f'qadm={rs.pression_sol_MPa} MPa — {rs.fondation.type.value}'],
        ['Durability', f'Exposure {"XS1" if rs.distance_mer_km<5 else "XC2"} — EN 206', f'Cover {40 if rs.distance_mer_km<5 else 30}mm'],
    ]
    _add_styled_table(doc, ['Domain', 'Standard', 'Value'], hyp_data,
                      col_widths=[3.8, 8.2, 5.0], header_color='43A956')

    doc.add_paragraph()  # Spacer

    # ── SECTION 3: LOAD TAKEDOWN — COLUMNS ────────────────────────
    doc.add_page_break()
    doc.add_heading('3. LOAD TAKEDOWN — COLUMNS (EC2/EC8)', level=1)

    desc_p = doc.add_paragraph(f'Spans {d.portee_min_m}–{d.portee_max_m} m — grid {d.nb_travees_x}×{d.nb_travees_y} — ULS combination: {1.35*rs.charge_G_kNm2+1.5*rs.charge_Q_kNm2:.1f} kN/m².')
    for run in desc_p.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    pot_headers = ['Level', 'NEd (kN)', 'Section', 'Bars', 'Ø (mm)', 'Tie Ø', 'Tie sp.', 'ρ (%)', 'NRd (kN)', 'NEd/NRd', 'Check']
    pot_rows = []
    for pt in rs.poteaux:
        pot_rows.append([
            pt.niveau,
            f'{pt.NEd_kN:.1f}',
            f'{pt.section_mm}×{pt.section_mm}',
            str(pt.nb_barres),
            str(pt.diametre_mm),
            str(pt.cadre_diam_mm),
            str(pt.espacement_cadres_mm),
            f'{pt.taux_armature_pct:.2f}',
            f'{pt.NRd_kN:.1f}',
            f'{pt.ratio_NEd_NRd:.2f}',
            '✓' if pt.verif_ok else '✗',
        ])

    _add_styled_table(doc, pot_headers, pot_rows,
                      col_widths=[1.8, 2.0, 2.4, 1.6, 1.6, 1.6, 2.0, 1.8, 2.0, 2.0, 1.2],
                      header_color='43A956', zebra=True)

    note_p = doc.add_paragraph('ρ = reinforcement ratio (EC2: 0.1% ≤ ρ ≤ 4%) | NEd/NRd < 1 required')
    for run in note_p.runs:
        run.font.size = Pt(8)
        run.font.italic = True

    doc.add_paragraph()  # Spacer

    # ── SECTION 4: BEAM DESIGN ──────────────────────────
    doc.add_page_break()
    doc.add_heading('4. BEAM DESIGN (EC2)', level=1)

    for pout in [rs.poutre_principale, rs.poutre_secondaire]:
        if pout is None:
            continue

        doc.add_heading(f'{pout.type} Beam — span {pout.portee_m} m', level=2)

        pout_data = [[
            str(pout.b_mm),
            str(pout.h_mm),
            f'{pout.As_inf_cm2:.1f}',
            f'{pout.As_sup_cm2:.1f}',
            f'HA{pout.etrier_diam_mm}',
            f'sp.={pout.etrier_esp_mm}mm',
            f'{pout.portee_m:.2f}',
        ]]
        _add_styled_table(doc, ['b (mm)', 'h (mm)', 'As bot (cm²)', 'As top (cm²)', 'Stirrups', 'Sp. stirrups', 'Span (m)'],
                          pout_data, col_widths=[2.2, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2], header_color='43A956')

        vf = '✓ OK' if pout.verif_fleche else '⚠ To verify'
        vt = '✓ OK' if pout.verif_effort_t else '⚠ To verify'
        vf_p = doc.add_paragraph(f'Deflection check: {vf} | Shear: {vt}')
        for run in vf_p.runs:
            run.font.size = Pt(8)
            run.font.italic = True

        doc.add_paragraph()  # Spacer

    # ── SECTION 5: SLAB DESIGN ────────────────────────────
    doc.add_heading('5. SLAB DESIGN (EC2)', level=1)

    dalle = rs.dalle
    dalle_data = [
        ['Thickness', f'{dalle.epaisseur_mm} mm', f'e ≥ L/35 = {rs.params.portee_min_m/35*1000:.0f} mm'],
        ['As x (cm²/m)', f'{dalle.As_x_cm2_ml:.1f}', 'Reinforcement along main span'],
        ['As y (cm²/m)', f'{dalle.As_y_cm2_ml:.1f}', 'Reinforcement along secondary span'],
        ['Allowable deflection', f'{dalle.fleche_admissible_mm:.1f} mm', f'L/250 = {rs.params.portee_min_m/250*1000:.1f} mm'],
        ['Verification', '✓ Compliant' if dalle.verif_ok else '⚠ To verify', ''],
    ]
    _add_styled_table(doc, ['Parameter', 'Value', 'Justification'], dalle_data,
                      col_widths=[4.2, 3.3, 7.5], header_color='43A956')

    doc.add_paragraph()  # Spacer

    # ── SECTION 6: PARTITIONS AND MASONRY ───────────────────────
    doc.add_page_break()
    doc.add_heading('6. PARTITIONS AND MASONRY', level=1)

    cl = rs.cloisons
    surf_p = doc.add_paragraph(f'Total partition area estimated: {int(cl.surface_totale_m2)} m² (load-bearing {int(cl.surface_separative_m2)} m² | light {int(cl.surface_legere_m2)} m² | ducts {int(cl.surface_gaines_m2)} m²)')
    for run in surf_p.runs:
        run.font.size = Pt(9)

    rec_p = doc.add_paragraph(f'Recommended option: {cl.option_recommandee.value} — load retained: {cl.charge_dalle_kn_m2} kN/m²')
    for run in rec_p.runs:
        run.font.size = Pt(9)
        run.font.bold = True

    doc.add_paragraph()  # Spacer

    # Partition options table
    cl_headers = ['Option', 'Th. (cm)', 'Load (kN/m²)', 'Unit price (FCFA/m²)', 'Recommended', 'Main advantages']
    cl_rows = []
    for opt in cl.options:
        est_rec = opt.type == cl.option_recommandee
        rec_text = '★ Recommended' if est_rec else '—'
        cl_rows.append([
            opt.materiau[:35],
            str(opt.epaisseur_cm),
            str(opt.charge_kn_m2),
            _format_number(opt.prix_fcfa_m2),
            rec_text,
            ' | '.join(opt.avantages[:2]),
        ])

    _add_styled_table(doc, cl_headers, cl_rows,
                      col_widths=[3.8, 1.6, 2.4, 2.4, 2.4, 7.4], header_color='43A956', zebra=True)

    adv_p = doc.add_paragraph('ℹ If multiple options have significant cost impacts, submit them to the owner before validation.')
    adv_p.paragraph_format.left_indent = Cm(0.5)
    for run in adv_p.runs:
        run.font.size = Pt(8)
        run.font.italic = True

    doc.add_paragraph()  # Spacer

    # ── SECTION 7: FOUNDATIONS ───────────────────────────────────
    doc.add_page_break()
    doc.add_heading('7. FOUNDATION DESIGN (EC7 + DTU 13.2)', level=1)

    fond = rs.fondation
    just_p = doc.add_paragraph(f'Justification: {fond.justification}')
    for run in just_p.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    doc.add_paragraph()  # Spacer

    fond_data = [
        ['Type', fond.type.value, 'Adapted to soil conditions and building height'],
    ]

    if fond.nb_pieux > 0:
        fond_data += [
            ['Pile diameter', f'Ø{fond.diam_pieu_mm} mm', 'Drilled with hollow auger'],
            ['Pile length', f'{fond.longueur_pieu_m:.1f} m', 'To competent stratum'],
            ['Reinforcement', f'As = {fond.As_cm2:.1f} cm²', 'HA500B cage full length'],
            ['Total piles', str(fond.nb_pieux), 'Estimated — to be confirmed by geotechnical engineer'],
        ]
    else:
        fond_data += [
            ['Footing width', f'{fond.largeur_semelle_m:.2f} m', 'Square section'],
            ['Depth', f'{fond.profondeur_m:.1f} m', 'Below frost depth + competent stratum'],
        ]

    _add_styled_table(doc, ['Parameter', 'Value', 'Remark'], fond_data,
                      col_widths=[4.2, 3.3, 7.5], header_color='43A956')

    # Foundation cost note
    if fond.nb_pieux > 0:
        try:
            from prix_marche import get_prix_structure
            px = get_prix_structure(rs.params.ville)
            cout_pieux = fond.nb_pieux * fond.longueur_pieu_m * px.pieu_fore_d800_ml

            impact_p = doc.add_paragraph(f'ℹ Foundation cost impact: {_format_fcfa(cout_pieux)} estimated ({cout_pieux/boq.total_bas_fcfa*100:.0f}% of structural budget). Deep foundations = highest cost item after main structure.')
            impact_p.paragraph_format.left_indent = Cm(0.5)
            for run in impact_p.runs:
                run.font.size = Pt(8)
                run.font.italic = True
        except Exception as e:
            logger.warning(f"Foundation cost calculation failed: {e}")

    doc.add_paragraph()  # Spacer

    # ── SECTION 8: SEISMIC ANALYSIS ─────────────────────────────
    doc.add_page_break()
    doc.add_heading('8. SEISMIC ANALYSIS (EC8 — NF EN 1998-1)', level=1)

    sism = rs.sismique
    sism_p = doc.add_paragraph(f'Seismic zone {sism.zone} — ag = {sism.ag_g}g — T₁ = {sism.T1_s}s — Fb = {sism.Fb_kN:.0f} kN')
    for run in sism_p.runs:
        run.font.size = Pt(9)
        run.font.italic = True

    doc.add_paragraph()  # Spacer

    sism_data = [
        ['Acceleration ag', f'{sism.ag_g}g = {sism.ag_g*9.81:.2f} m/s²', 'National annex'],
        ['Soil factor S', '1.15', 'Soil type C — EC8 Table 3.2'],
        ['q factor', '1.5 (DCL)', 'EC8 §5.2.2.2'],
        ['Period T₁', f'{sism.T1_s} s', 'EC8 §4.3.3.2 — approximate method'],
        ['Base shear Fb', f'{sism.Fb_kN:.0f} kN', 'Fb = Sd(T₁) × m × λ'],
        ['DCL compliance', '✓ Compliant' if sism.conforme_DCL else '⚠ Further analysis needed', ''],
    ]
    _add_styled_table(doc, ['Parameter', 'Value', 'Reference'], sism_data,
                      col_widths=[5.25, 5.25, 5.5], header_color='43A956')

    doc.add_paragraph()  # Spacer

    # Seismic provisions
    for disp in sism.dispositions:
        prefix = '⚠' if '⚠' in disp else '•'
        disp_p = doc.add_paragraph(f'{prefix} {disp.replace("⚠ ", "")}')
        for run in disp_p.runs:
            run.font.size = Pt(9)
            if '⚠' in disp:
                run.font.bold = True
                run.font.color.rgb = ORANGE

    doc.add_paragraph()  # Spacer

    # ── SECTION 9: ANALYSIS AND RECOMMENDATIONS ───────────────────────
    doc.add_page_break()
    doc.add_heading('9. ANALYSIS AND RECOMMENDATIONS', level=1)

    # Engineer notes
    doc.add_heading('Summary:', level=2)
    note_p = doc.add_paragraph(ana.note_ingenieur)
    for run in note_p.runs:
        run.font.color.rgb = NAVY
        run.font.italic = True

    doc.add_paragraph()  # Spacer

    # Strengths / Alerts
    if ana.points_forts or ana.alertes:
        # Create a 2-column table for strengths and alerts
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Column 1: Strengths
        col1_cell = table.rows[0].cells[0]
        col1_cell.width = Cm(8)
        col1_para = col1_cell.paragraphs[0]
        col1_para.text = ''

        if ana.points_forts:
            heading = col1_cell.add_heading('✅ Project strengths', level=2)
            for f in ana.points_forts:
                p = col1_cell.add_paragraph(f'• {f}')
                for run in p.runs:
                    run.font.size = Pt(9)

        # Column 2: Alerts
        col2_cell = table.rows[0].cells[1]
        col2_cell.width = Cm(8)
        col2_para = col2_cell.paragraphs[0]
        col2_para.text = ''

        if ana.alertes:
            heading = col2_cell.add_heading('⚠ Points of attention', level=2)
            for a in ana.alertes:
                p = col2_cell.add_paragraph(f'• {a}')
                for run in p.runs:
                    run.font.size = Pt(9)
                    run.font.color.rgb = ORANGE

        doc.add_paragraph()  # Spacer

    # Recommendations
    if ana.recommandations:
        doc.add_heading('Recommendations:', level=2)

        rec_headers = ['No.', 'Recommendation']
        rec_rows = []
        for i, r in enumerate(ana.recommandations):
            rec_rows.append([str(i+1), r])

        _add_styled_table(doc, rec_headers, rec_rows,
                          col_widths=[1.2, 16.8], header_color='43A956')

    doc.add_paragraph()  # Spacer

    # ── SECTION 10: BOQ — STRUCTURE ───────────────────────────
    doc.add_page_break()
    doc.add_heading('10. BILL OF QUANTITIES AND PRICING — STRUCTURE', level=1)

    boq_desc = doc.add_paragraph(f'Unit prices market {rs.params.ville} 2026 (supply and installation). Estimated margin ±15%.')
    for run in boq_desc.runs:
        run.font.size = Pt(8)
        run.font.italic = True

    doc.add_paragraph()  # Spacer

    # BOQ Table
    boq_headers = ['Item', 'Description', 'Qty', 'Unit', 'Unit price (FCFA)', 'Low amount', 'High amount']
    boq_rows = [
        ['1', 'Excavation — stripping + mechanical excavation', _format_number(boq.terrassement_m3), 'm³', '8 500', _format_fcfa(boq.cout_terr_fcfa), _format_fcfa(int(boq.cout_terr_fcfa*1.10))],
        ['2', 'Foundations — piles/footings/raft reinforced concrete', '—', 'lump sum', '—', _format_fcfa(boq.cout_fond_fcfa), _format_fcfa(int(boq.cout_fond_fcfa*1.20))],
        ['3a', f'Concrete {rs.classe_beton} BPE — structure ({_format_number(boq.beton_structure_m3,0)} m³)', _format_number(boq.beton_structure_m3,0), 'm³', '185 000', _format_fcfa(boq.cout_beton_fcfa), _format_fcfa(int(boq.cout_beton_fcfa*1.10))],
        ['3b', f'Steel {rs.classe_acier} supply and install ({_format_number(boq.acier_kg,0)} kg)', _format_number(boq.acier_kg,0), 'kg', '810', _format_fcfa(boq.cout_acier_fcfa), _format_fcfa(int(boq.cout_acier_fcfa*1.10))],
        ['3c', f'Formwork all faces ({_format_number(boq.coffrage_m2,0)} m²)', _format_number(boq.coffrage_m2,0), 'm²', '18 000', _format_fcfa(boq.cout_coffrage_fcfa), _format_fcfa(int(boq.cout_coffrage_fcfa*1.10))],
        ['4', 'Masonry — 15cm blocks plastered 2 faces', _format_number(boq.maconnerie_m2,0), 'm²', '24 000', _format_fcfa(boq.cout_maco_fcfa), _format_fcfa(int(boq.cout_maco_fcfa*1.15))],
        ['5', f'Waterproofing flat roof ({_format_number(boq.etancheite_m2,0)} m²)', _format_number(boq.etancheite_m2,0), 'm²', '18 500', _format_fcfa(boq.cout_etanch_fcfa), _format_fcfa(int(boq.cout_etanch_fcfa*1.10))],
        ['6', 'Miscellaneous — joints, parapets, openings', '—', 'lump sum', '—', _format_fcfa(boq.cout_divers_fcfa), _format_fcfa(int(boq.cout_divers_fcfa*1.10))],
    ]

    _add_styled_table(doc, boq_headers, boq_rows,
                      col_widths=[1.2, 7.0, 1.8, 1.4, 2.6, 2.8, 3.0], header_color='43A956', zebra=True)

    # BOQ Total row
    total_table = doc.add_table(rows=1, cols=7)
    total_table.style = 'Table Grid'

    total_row = total_table.rows[0]
    total_cells = [
        '', '', '', '',
        'TOTAL STRUCTURE',
        _format_fcfa(boq.total_bas_fcfa),
        _format_fcfa(boq.total_haut_fcfa)
    ]

    for i, cell_text in enumerate(total_cells):
        cell = total_row.cells[i]
        cell.text = str(cell_text)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT if i >= 4 else WD_ALIGN_PARAGRAPH.LEFT
        run = p.runs[0]
        run.font.bold = True
        run.font.size = Pt(9)
        _set_cell_background(cell, 'E8F5EA')

    doc.add_paragraph()  # Spacer

    # Ratios table
    doc.add_heading('Structural ratios', level=2)

    ratios_headers = ['Indicator', 'Low value', 'High value', 'Note']
    ratios_rows = [
        ['Total built area', _format_number(boq.surface_batie_m2, 0, 'm²'), '—', f'Footprint {int(d.surface_emprise_m2)} m² × {d.nb_niveaux} levels'],
        ['Cost / m² built', f'{boq.ratio_fcfa_m2_bati:,} FCFA/m²'.replace(',', ' '), f'{int(boq.ratio_fcfa_m2_bati*1.15):,} FCFA/m²'.replace(',', ' '), 'Structure only — excludes MEP, finishes, site works'],
        ['Cost / m² habitable', f'{boq.ratio_fcfa_m2_habitable:,} FCFA/m²'.replace(',', ' '), f'{int(boq.ratio_fcfa_m2_habitable*1.15):,} FCFA/m²'.replace(',', ' '), 'Habitable area ≈ 78% of built area'],
        ['TOTAL STRUCTURE COST', _format_fcfa(boq.total_bas_fcfa), _format_fcfa(boq.total_haut_fcfa), 'Estimation ±15%'],
    ]

    _add_styled_table(doc, ratios_headers, ratios_rows,
                      col_widths=[4.8, 3.0, 3.0, 4.2], header_color='43A956')

    doc.add_paragraph()  # Spacer

    # ── FOOTER ──────────────────────────────────────────────────
    footer_p = doc.add_paragraph('\n\nDocument generated by Tijan AI — tijan.ai')
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(8)
        run.font.italic = True
        run.font.color.rgb = GRIS

    # Save to bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()
